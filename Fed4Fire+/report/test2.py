#!/usr/bin/env python3
from docx import Document
import requests
from copy import deepcopy
from shutil import copyfile, rmtree

def get_para_data(output_doc_name, paragraph):
    """
    Write the run to the new file and then set its font, bold, alignment, color etc. data.
    """

    output_para = output_doc_name.add_paragraph()
    for run in paragraph.runs:
        output_run = output_para.add_run(run.text)
        # Run's bold data
        output_run.bold = run.bold
        # Run's italic data
        output_run.italic = run.italic
        # Run's underline data
        output_run.underline = run.underline
        # Run's color data
        output_run.font.color.rgb = run.font.color.rgb
        # Run's font data
        output_run.style.name = run.style.name
    # Paragraph's alignment data
    output_para.paragraph_format.alignment = paragraph.paragraph_format.alignment

copyfile("test.docx","test2.docx")

document = Document('test2.docx')

table = document.tables[0]

sessions = [(153,"name",1)]
interface = "131.114.72.76:8080"

row_idx = 2
for session,name,moment in sessions:
    r = requests.get(f"http://{interface}/testbed/{session}/accuracy")
    errors = r.json()["data"]
    print(errors)
    r = requests.get(f"http://{interface}/testbed/{session}/footprint")
    footprint = r.json()["data"]
    print(footprint)

    cells = table.rows[row_idx].cells
    cells[0].paragraphs[0].runs[0].text = name
    cells[1].paragraphs[0].runs[0].text = f'{errors[moment]["L"]["intra"]["mean"]*100:.0f}%'
    cells[2].paragraphs[0].runs[0].text = f'{errors[moment]["B"]["intra"]["mean"]*100:.0f}%'
    cells[3].paragraphs[0].runs[0].text = f'{errors[moment]["L"]["inter"]["mean"]*100:.0f}%'
    cells[4].paragraphs[0].runs[0].text = f'{errors[moment]["B"]["inter"]["mean"]*100:.0f}%'

    cells[5].paragraphs[0].runs[0].text = f'{footprint[moment]["cpu"]["mean"]:.2f}%'
    cells[6].paragraphs[0].runs[0].text = f'{footprint[moment]["mem"]["mean"]:.2f} MB'
    cells[7].paragraphs[0].runs[0].text = f'{footprint[moment]["tx"]["mean"]:.2f} KBps'
    cells[8].paragraphs[0].runs[0].text = f'{footprint[moment]["rx"]["mean"]:.2f} KBps'

    cells[9].paragraphs[0].runs[0].text = f'{errors[moment]["time"]:.0f} s'
    cells[9].paragraphs[1].runs[0].text = f'({errors[moment]["time"]//60:.0f} m {errors[moment]["time"]%60:.0f} s)'

    row_idx+=1




document.save('test2.docx')