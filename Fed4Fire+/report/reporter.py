#!/usr/bin/env python3
from docx import Document
import requests
from copy import deepcopy
from shutil import copyfile, rmtree
import re
from statistics import mean
documents = {
    # "links default 20 1.2": [(140,"exp1",1),(129,"exp2",1)],
    # "links reactive 20 1.2": [(151,"exp1",1),(8,"exp2",1)],

    "nodes 40 reactive": [(204,"init",0),(172,"exp1",1),(173,"exp2",1),(174,"exp3",1)],
    "leaders 40 reactive": [(175,"exp1",0),(173,"exp2",0),(176,"exp3",0)],
    "links 40 reactive": [(204,"exp1",1),(206,"exp2",1)],

    "nodes 40 default": [(201,"init",0),(199,"exp1",1),(200,"exp2",1),(201,"exp3",1)],
    "leaders 40 default": [(202,"exp1",0),(200,"exp2",0),(203,"exp3",0)],
    "links 40 default": [(207,"exp1",1),(208,"exp2",1)],

    "nodes 30 reactive": [(84,"init",0),(167,"exp1",1),(168,"exp2",1),(169,"exp3",1)],
    "leaders 30 reactive": [(170,"exp1",0),(168,"exp2",0),(171,"exp3",0)],
    "links 30 reactive": [(209,"exp1",1),(210,"exp2",1)],

    "nodes 30 default": [(184,"init",0),(184,"exp1",1),(185,"exp2",1),(186,"exp3",1)],
    "leaders 30 default": [(191,"exp1",0),(186,"exp2",0),(192,"exp3",0)],
    "links 30 default": [(211,"exp1",1),(212,"exp2",1)],

    "nodes 20 reactive": [(43,"init",0),(43,"exp1",1),(109,"exp2",1),(141,"exp3",1)],
    "leaders 20 reactive": [(142,"exp1",0),(109,"exp2",0),(144,"exp3",0)],
    "links 20 reactive": [(80,"exp1",1),(83,"exp2",1)],

    "nodes 20 default": [(155,"init",0),(160,"exp1",1),(146,"exp2",1),(155,"exp3",1)],
    "leaders 20 default": [(156,"exp1",0),(146,"exp2",0),(157,"exp3",0)],
    "links 20 default": [(145,"exp1",0),(159,"exp2",0)],
}

ext = [
    (0,"%"),
    (0,"%"),
    (0,"%"),
    (0,"%"),
    (2,"%"),
    (2," MB"),
    (2," KBps"),
    (2," KBps"),
    (0," s"),
]

def set_text(cell,text, paragraph=0, run=0):
    cell.paragraphs[paragraph].runs[run].text = text

def get_text(cell,paragraph=0, run=0):
    text = cell.paragraphs[paragraph].runs[run].text
    return float(re.search(r'[\d.]+', text).group())

for doc,sessions in documents.items():
    copyfile("test.docx",f"{doc}.docx")

    document = Document('test2.docx')

    table = document.tables[0]

    interface = "131.114.72.76:8080"

    row_idx = 2
    for session,name,moment in sessions:
        if session is not None:
            r = requests.get(f"http://{interface}/testbed/{session}/accuracy")
            errors = r.json()["data"]
            print(errors)
            r = requests.get(f"http://{interface}/testbed/{session}/footprint")
            footprint = r.json()["data"]
            print(footprint)
        else:
            errors = [{"L":{"intra":{"mean":0}},"B":{"intra":{"mean":0}}} for m in range(moment)]
            footprint = [{"cpu":{"mean":0},"mem":{"mean":0},"tx":{"mean":0},"rx":{"mean":0},"time":0} for m in range(moment)]

        cells = table.rows[row_idx].cells
        set_text(cells[0],name)
        set_text(cells[1], f'{errors[moment]["L"]["intra"]["mean"]*100:.0f}%')
        set_text(cells[2], f'{errors[moment]["B"]["intra"]["mean"]*100:.0f}%')
        set_text(cells[3], f'{errors[moment]["L"]["inter"]["mean"]*100:.0f}%')
        set_text(cells[4], f'{errors[moment]["B"]["inter"]["mean"]*100:.0f}%')

        set_text(cells[5], f'{footprint[moment]["cpu"]["mean"]:.2f}%')
        set_text(cells[6], f'{footprint[moment]["mem"]["mean"]:.2f} MB')
        set_text(cells[7], f'{footprint[moment]["tx"]["mean"]/1000:.2f} KBps')
        set_text(cells[8], f'{footprint[moment]["rx"]["mean"]/1000:.2f} KBps')

        set_text(cells[9], f'{errors[moment]["time"]:.0f} s')
        set_text(cells[9], f'({errors[moment]["time"]//60:.0f} m {errors[moment]["time"]%60:.0f} s)', paragraph=1)

        row_idx+=1

    def remove_row(table, row):
        tbl = table._tbl
        tr = row._tr
        tbl.remove(tr)
        
    rows = [table.rows[row_idx] for row_idx in range(row_idx,6)]
    for row in rows:
        remove_row(table,row)

    cells = table.rows[-1].cells
    set_text(cells[0],"Average")
    for i in range(1,10):
        dec,ext_ = ext[i-1]
        set_text(cells[i],f"{mean([get_text(row.cells[i]) for row in table.rows[2:-1]]):.{dec}f}{ext_}")

    document.save(f"{doc}.docx")

from zipfile import ZipFile

with ZipFile('file.zip', 'w') as zipObj2:
   # Add multiple files to the zip
   for doc,sessions in documents.items():
       zipObj2.write(f"{doc}.docx")